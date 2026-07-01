from docx import Document


def export_report(report, output_file):

    doc = Document()

    doc.add_heading(
        report["title"],
        level=0
    )

    for section in report["sections"]:

        doc.add_heading(
            section["heading"],
            level=1
        )

        doc.add_paragraph(
            section["content"]
        )

    doc.save(output_file)